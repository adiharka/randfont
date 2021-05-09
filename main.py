# Created by AdiHarka
# Feel free to use it, i know you are too lazy writing some paper
# This code make new fresh docx file with randomized font, but doesn't copy it's format (like font size, bold, italic)
# So make sure to edit your document after converting, for human peace no waste of time

from docx import Document
import random
import os

# List of font
fontlist = ['Times New Roman', 'Calibri', 'And other font you want to include']

# Read file (make sure your path is right)
file = os.path.abspath('yourfile.docx')
print(file)
doc = Document(path)

# Convert docx to String
data = []
for i in doc.paragraphs:
    data.append(i.text)

# Make docx from String
doc_res = Document()
for i in data:
    paragraph = doc_res.add_paragraph()
    for j in i:
        run = paragraph.add_run(j)
        run.font.name = random.choice(fontlist)

# Save new docx file
doc_res.save('resultfile.docx')